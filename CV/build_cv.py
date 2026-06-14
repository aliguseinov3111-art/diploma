from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# Styles
style_normal = doc.styles['Normal']
style_normal.font.name = 'Calibri'
style_normal.font.size = Pt(10.5)

def add_name(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x1F, 0x39, 0x64)
    p.space_after = Pt(2)
    return p

def add_subtitle(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    p.space_after = Pt(2)
    return p

def add_contacts(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(10)
    p.space_after = Pt(6)
    return p

def add_positioning(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor(0x1F, 0x39, 0x64)
    p.space_after = Pt(8)
    return p

def add_section_header(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor(0x1F, 0x39, 0x64)
    p.space_before = Pt(10)
    p.space_after = Pt(2)
    # Bottom border
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1F3964')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def add_employer(doc, company, location, period, subtitle=None):
    p = doc.add_paragraph()
    p.space_before = Pt(8)
    p.space_after = Pt(0)
    run = p.add_run(f"{company}, {location}")
    run.bold = True
    run.font.size = Pt(10.5)
    run2 = p.add_run(f"  |  {period}")
    run2.font.size = Pt(10)
    run2.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    if subtitle:
        p2 = doc.add_paragraph()
        p2.space_before = Pt(0)
        p2.space_after = Pt(1)
        r = p2.add_run(subtitle)
        r.font.size = Pt(9.5)
        r.italic = True
        r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

def add_role(doc, title, period, note=None):
    p = doc.add_paragraph()
    p.space_before = Pt(5)
    p.space_after = Pt(0)
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(10.5)
    run2 = p.add_run(f"  |  {period}")
    run2.font.size = Pt(10)
    run2.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    if note:
        p2 = doc.add_paragraph()
        p2.space_before = Pt(0)
        p2.space_after = Pt(1)
        r = p2.add_run(f"({note})")
        r.font.size = Pt(9.5)
        r.italic = True
        r.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

def add_bullet(doc, text, bold_part=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(1)
    if bold_part and text.startswith(bold_part):
        run1 = p.add_run(bold_part)
        run1.bold = True
        run1.font.size = Pt(10.5)
        run2 = p.add_run(text[len(bold_part):])
        run2.font.size = Pt(10.5)
    else:
        run = p.add_run(text)
        run.font.size = Pt(10.5)
    return p

def add_plain(doc, text, bold=False, size=10.5):
    p = doc.add_paragraph()
    p.space_before = Pt(2)
    p.space_after = Pt(2)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    return p

def add_edu_entry(doc, university, degree, years):
    p = doc.add_paragraph()
    p.space_before = Pt(5)
    p.space_after = Pt(0)
    run = p.add_run(university)
    run.bold = True
    run.font.size = Pt(10.5)
    p2 = doc.add_paragraph()
    p2.space_before = Pt(0)
    p2.space_after = Pt(0)
    run2 = p2.add_run(degree)
    run2.font.size = Pt(10.5)
    p3 = doc.add_paragraph()
    p3.space_before = Pt(0)
    p3.space_after = Pt(3)
    run3 = p3.add_run(years)
    run3.font.size = Pt(10)
    run3.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

# ─── HEADER ───────────────────────────────────────────────────────────────────
# Header table: photo on right, name/contacts on left
from docx.oxml.ns import nsmap as _nsmap
from lxml import etree

header_tbl = doc.add_table(rows=1, cols=2)
header_tbl.style = 'Table Grid'
# Remove borders from table
tbl_pr_el = header_tbl._tbl.find(qn('w:tblPr'))
if tbl_pr_el is None:
    tbl_pr_el = OxmlElement('w:tblPr')
    header_tbl._tbl.insert(0, tbl_pr_el)
tbl_borders = OxmlElement('w:tblBorders')
for border_name in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
    b = OxmlElement(f'w:{border_name}')
    b.set(qn('w:val'), 'none')
    b.set(qn('w:sz'), '0')
    b.set(qn('w:space'), '0')
    b.set(qn('w:color'), 'auto')
    tbl_borders.append(b)
tbl_pr_el.append(tbl_borders)

# Set column widths: text ~13cm, photo ~3.5cm
from docx.oxml import OxmlElement as OE
tblGrid = OxmlElement('w:tblGrid')
for w_twips in (7400, 1980):
    gc = OxmlElement('w:gridCol')
    gc.set(qn('w:w'), str(w_twips))
    tblGrid.append(gc)
header_tbl._tbl.insert(1, tblGrid)

left_cell = header_tbl.cell(0, 0)
right_cell = header_tbl.cell(0, 1)

# Set vertical alignment top for both cells
for cell in (left_cell, right_cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    v_align = OxmlElement('w:vAlign')
    v_align.set(qn('w:val'), 'top')
    tc_pr.append(v_align)

# Left cell: name + city + contacts
p_name = left_cell.paragraphs[0]
p_name.alignment = WD_ALIGN_PARAGRAPH.LEFT
r_name = p_name.add_run("Гусейнов Али")
r_name.bold = True
r_name.font.size = Pt(18)
r_name.font.color.rgb = RGBColor(0x1F, 0x39, 0x64)
p_name.space_after = Pt(4)

p_city = left_cell.add_paragraph()
p_city.alignment = WD_ALIGN_PARAGRAPH.LEFT
r_city = p_city.add_run("Москва, Россия")
r_city.font.size = Pt(10)
r_city.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
p_city.space_after = Pt(2)

p_contacts = left_cell.add_paragraph()
p_contacts.alignment = WD_ALIGN_PARAGRAPH.LEFT
r_contacts = p_contacts.add_run("+7 (984) 282-65-57  ·  aliguseinov03@mail.ru")
r_contacts.font.size = Pt(10)
p_contacts.space_after = Pt(4)

# Right cell: photo
right_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
run_photo = right_cell.paragraphs[0].add_run()
run_photo.add_picture(r"d:\diploma\CV\фото.jpg", width=Cm(3.2))

# ─── ОБРАЗОВАНИЕ ──────────────────────────────────────────────────────────────
add_section_header(doc, "Образование")

add_edu_entry(
    doc,
    "НИУ МГСУ (Национальный исследовательский Московский государственный строительный университет)",
    "Магистратура — 08.04.01 Строительство, профиль: Промышленное и гражданское строительство",
    "2024 – 2026"
)
add_edu_entry(
    doc,
    "НИУ МГСУ",
    "Профессиональная переподготовка — «Основы организации и технологий строительного производства»",
    "Апрель – Сентябрь 2025"
)
add_edu_entry(
    doc,
    "Государственный Университет Управления",
    "Бакалавриат — Экономика и управление инвестиционно-строительной деятельностью",
    "2020 – 2024"
)

# ─── ОПЫТ РАБОТЫ ──────────────────────────────────────────────────────────────
add_section_header(doc, "Профессиональный опыт")

# ВЕРНИСАЖ ГРУПП
add_employer(
    doc,
    "Вернисаж Групп", "Москва", "Июль 2024 — Июль 2026",
    "Подрядчик по отделочным, фасадным и инженерным работам  |  Коммерческая и жилая недвижимость"
)

add_role(doc, "Руководитель проекта", "Июль 2025 — Июль 2026",
         "повышен с позиции инженера через 12 месяцев")
add_bullet(doc, "Управлял полным бюджетом проекта: калькулирование расценок, план-фактный анализ, контроль отклонений от сметных показателей")
add_bullet(doc, "Руководил командой проекта: прорабы, инженер ПТО, специалисты по охране труда, субподрядчики и поставщики")
add_bullet(doc, "Согласовывал дополнительные работы и изменения сметных цен с заказчиком; подготавливал и подписывал дополнительные соглашения к договорам")
add_bullet(doc, "Вёл претензионную переписку с заказчиками и подрядчиками; сопровождал урегулирование спорных вопросов по объёмам и стоимости")
add_bullet(doc, "Координировал поставки материалов: анализировал рыночные предложения, вёл переговоры и подписывал договоры с поставщиками")
add_bullet(doc, "Обеспечивал подготовку исполнительной документации и соблюдение требований охраны труда на объектах")
add_bullet(doc, "Проводил совещания с заказчиком; принимал оперативные решения в рамках бюджета и сроков")

add_role(doc, "Инженер технического офиса (контроль и планирование затрат)", "Июль 2024 — Июль 2025")
add_bullet(doc, "Калькулировал бюджеты и расценки на отделочные и инженерные работы на стадиях от коммерческого предложения до рабочей документации")
add_bullet(doc, "Участвовал в реализации проектов на объектах культурного наследия (ОКН) — работа в условиях особых регуляторных и технических требований")
add_bullet(doc, "Формировал тендерную документацию; анализировал коммерческие предложения поставщиков и субподрядчиков для выбора оптимальных ценовых решений")
add_bullet(doc, "Вёл план-фактный анализ бюджетов: выявлял отклонения фактических затрат от сметных, инициировал корректирующие меры")
add_bullet(doc, "Верифицировал акты КС-2 и справки КС-3 на соответствие договорным объёмам и расценкам — как на стороне подрядчика, так и заказчика")
add_bullet(doc, "Согласовывал стоимостные параметры с проектировщиками и заказчиком; подготавливал аналитические отчёты по стоимостным показателям для руководства; вёл управленческий учёт в 1С")
add_bullet(doc, "Контролировал исполнение подрядных договоров по срокам, объёмам и порядку оплат")

# ГЭС КОНСТРАКШН
add_employer(
    doc,
    "ГЭС Констракшн", "Москва", "Декабрь 2022 — Март 2024",
    "Генеральный подрядчик, полный цикл строительства  |  Проект: ЖК ДримТауэрс"
)

add_role(doc, "Ассистент технического отдела", "Май 2023 — Март 2024",
         "повышен с позиции стажёра через 5 месяцев")
add_bullet(doc, "Формировал коммерческие сметы и рассчитывал удельные показатели стоимости работ и материалов на объекте ЖК ДримТауэрс")
add_bullet(doc, "Производил подсчёт объёмов работ по проектной и рабочей документации")
add_bullet(doc, "Вёл анализ расходов по проекту: контролировал соответствие фактических затрат плановым показателям")
add_bullet(doc, "Верифицировал акты КС-2 и справки КС-3 на соответствие договорным объёмам и расценкам")
add_bullet(doc, "Разрабатывал графики производства работ и графики движения рабочей силы")
add_bullet(doc, "Вёл претензионную переписку с подрядчиками и поставщиками")
add_bullet(doc, "Координировал работу со смежными и финансовым отделами в части стоимостных показателей проекта")
add_bullet(doc, "Контролировал исполнение подрядных договоров по срокам, объёмам и оплатам")

add_role(doc, "Стажёр технического отдела", "Декабрь 2022 — Май 2023")
add_bullet(doc, "Участвовал в формировании коммерческих смет и подсчёте объёмов работ по проектной документации")
add_bullet(doc, "Осваивал процедуры проверки актов КС-2/КС-3 и контроля договорной документации")
add_bullet(doc, "Поддерживал разработку графиков производства работ и взаимодействие с подрядчиками")

# ─── НАВЫКИ ───────────────────────────────────────────────────────────────────
add_section_header(doc, "Навыки")

p = doc.add_paragraph()
p.space_before = Pt(4)
p.space_after = Pt(1)
r = p.add_run("Программное обеспечение:  ")
r.bold = True
r.font.size = Pt(10.5)
r2 = p.add_run("MS Excel (продвинутый)  ·  1С Управление (уверенный)  ·  MS Project (средний)  ·  AutoCAD (средний)  ·  Revit (подсчёт объёмов)")
r2.font.size = Pt(10.5)

p2 = doc.add_paragraph()
p2.space_before = Pt(4)
p2.space_after = Pt(2)
r3 = p2.add_run("Профессиональные навыки:")
r3.bold = True
r3.font.size = Pt(10.5)

skills = [
    "Коммерческое ценообразование в строительстве; калькулирование бюджетов и расценок на разных стадиях проекта",
    "План-фактный анализ: выявление отклонений фактических затрат от плановых, подготовка аналитических отчётов",
    "Мониторинг рыночных цен на работы, материалы и услуги; анализ коммерческих предложений поставщиков и подрядчиков",
    "Подсчёт объёмов работ по проектной и рабочей документации",
    "Формирование тендерной документации; сравнительный анализ предложений",
    "Работа с договорной документацией: договоры, дополнительные соглашения, претензионная переписка",
    "Разработка графиков производства работ и движения рабочей силы",
    "Управление проектом полного цикла: бюджет, команда, сроки, исполнительная документация",
]
for s in skills:
    add_bullet(doc, s)

# ─── ЯЗЫКИ ────────────────────────────────────────────────────────────────────
add_section_header(doc, "Языки")
add_bullet(doc, "Русский — родной")
add_bullet(doc, "Английский — продвинутый уровень (чтение, письмо, устная речь)")
add_bullet(doc, "Азербайджанский — родной")
add_bullet(doc, "Турецкий — средний уровень")

# ─── ПЕРСОНАЛЬНЫЕ ДАННЫЕ ──────────────────────────────────────────────────────
add_section_header(doc, "Персональные данные")
p = doc.add_paragraph()
p.space_before = Pt(4)
p.space_after = Pt(2)
run = p.add_run("Дата рождения: 03.11.2001  ·  Семейное положение: холост  ·  Водительское удостоверение: кат. В")
run.font.size = Pt(10.5)

doc.save(r"d:\diploma\CV\CV Guseinov Ali (updated).docx")
print("Done")
