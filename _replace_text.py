"""
Replace text paragraphs in the original docx with rewritten content,
preserving all images, tables, headings, and captions in place.
"""
import re
import sys
import copy
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

sys.stdout.reconfigure(encoding='utf-8')

# ── Load ──────────────────────────────────────────────────────────────────────
doc = Document('02 ВКР Али.docx')

with open('02_VKR_Ali_rewrite_intro.md', 'r', encoding='utf-8') as f:
    rewrite_lines = f.readlines()

# ── Para helpers ──────────────────────────────────────────────────────────────
def has_image(para):
    return para._element.find('.//' + qn('a:blip')) is not None

def is_heading(para):
    return para.style.name.startswith('Heading')

def is_empty(para):
    return para.text.strip() == ''

def should_keep(para):
    """Paragraphs that must not be deleted."""
    if is_heading(para):   return True
    if has_image(para):    return True
    if is_empty(para):     return True
    text = para.text.strip()
    # Keep figure captions and table captions / notes
    if re.match(r'^Рис[\.\s]', text):  return True
    if text.startswith('Рисунок'):     return True
    if text.startswith('Таблица'):     return True
    if text.startswith('Примечание'):  return True
    if text.startswith('Table'):       return True
    return False

# ── Find a template Normal paragraph (for cloning style) ─────────────────────
template_para = None
for p in doc.paragraphs:
    if (p.style.name == 'Normal'
            and p.text.strip()
            and not has_image(p)
            and not re.match(r'^Рис[\.\s]', p.text.strip())):
        template_para = p
        break

def make_para_elem(text, template):
    """Clone template paragraph, wipe runs, insert new text run."""
    new_p = copy.deepcopy(template._element)
    # Remove everything that carries text content
    for tag in (qn('w:r'), qn('w:hyperlink'), qn('w:bookmarkStart'), qn('w:bookmarkEnd')):
        for child in new_p.findall(tag):
            new_p.remove(child)
    # Strip run-props stored inside pPr (keep pStyle/spacing/indent)
    pPr = new_p.find(qn('w:pPr'))
    if pPr is not None:
        for rPr in pPr.findall(qn('w:rPr')):
            pPr.remove(rPr)

    # Build run
    r = OxmlElement('w:r')
    # Copy run properties from first run of template if available
    tmpl_runs = template._element.findall(qn('w:r'))
    if tmpl_runs:
        rPr = tmpl_runs[0].find(qn('w:rPr'))
        if rPr is not None:
            r.append(copy.deepcopy(rPr))
    t_elem = OxmlElement('w:t')
    t_elem.text = text
    if text and (text[0] == ' ' or text[-1] == ' '):
        t_elem.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    r.append(t_elem)
    new_p.append(r)
    return new_p

# ── Parse rewrite markdown ────────────────────────────────────────────────────
FIG_RE = re.compile(r'^>?\s*Рис[\.т]?[\s\d\(]', re.IGNORECASE)
TABLE_RE = re.compile(r'^>?\s*Таблица\s')

def clean_md_line(raw):
    s = raw.strip()
    s = re.sub(r'^>\s*', '', s)                        # blockquote
    s = re.sub(r'^\d+[\.\)]\s+', '', s)               # numbered list
    s = re.sub(r'^[-*]\s+', '', s)                     # bullet
    s = re.sub(r'\[([^\]]+)\]\{\.underline\}', r'\1', s)  # underline span
    s = re.sub(r'\*\*([^*]+)\*\*', r'\1', s)          # bold
    s = re.sub(r'\*([^*]+)\*', r'\1', s)              # italic
    s = s.replace('\\*', '*').replace('\\.', '.').replace('\\-', '-')
    s = re.sub(r'<!--.*?-->', '', s)                   # html comments
    # Un-escape asterisks used in norms e.g. СНиП II-23-81\*
    s = s.replace('\\*', '*')
    return s.strip()

def parse_rewrite_sections(lines):
    sections = {}       # key → list[str]
    current_key = None
    current_lines = []

    for raw in lines:
        stripped = raw.rstrip('\n').strip()

        if stripped == '---':
            continue

        # Heading line
        if re.match(r'^#{1,4}\s', stripped):
            heading_text = stripped.lstrip('#').strip()
            # Skip meta-headings at the top of the file
            if ('ПЕРЕРАБОТАННЫЕ' in heading_text
                    or 'таблицы, формулы' in heading_text):
                continue
            # Save previous
            if current_key is not None:
                sections[current_key] = [l for l in current_lines if l]
            current_key = heading_text
            current_lines = []
            continue

        if current_key is None:
            continue

        # Skip placeholder markers
        if re.match(r'^\*\(', stripped) and stripped.endswith(')*'):
            continue
        if 'сохраняется без изменений' in stripped:
            continue
        # Skip figure references — they're already in original as captions
        if FIG_RE.match(stripped):
            continue
        # Skip table caption references — kept from original
        if TABLE_RE.match(stripped):
            continue

        if not stripped:
            continue

        cleaned = clean_md_line(stripped)
        if cleaned:
            current_lines.append(cleaned)

    if current_key is not None:
        sections[current_key] = [l for l in current_lines if l]

    return sections

rewrite_sections = parse_rewrite_sections(rewrite_lines)
print(f"Parsed {len(rewrite_sections)} rewrite sections")
for k, v in rewrite_sections.items():
    print(f"  [{len(v):3d} lines] {k[:70]}")

# ── Map original headings → rewrite section keys ──────────────────────────────
# Each entry: (original heading keyword, rewrite section keyword)
# The script will find the first rewrite section whose key CONTAINS the given keyword.
REPLACE_MAP = [
    ('ВВЕДЕНИЕ',                         'ВВЕДЕНИЕ'),
    ('1.1.',                             '1.1.'),
    ('1.2',                              '1.2'),
    ('1.3.',                             '1.3.'),
    ('1.4.',                             '1.4.'),
    ('1.5.',                             '1.5.'),
    ('1.6.',                             '1.6.'),
    # 1.7 → SKIP (static calculation tables — keep original)
    ('1.8.',                             '1.8.'),
    ('1.9.',                             '1.9.'),
    ('2.1.',                             '2.1.'),
    ('2.2.',                             '2.2.'),
    ('2.3. Совершенствование',           '2.3. Совершенствование'),
    ('2.3. Алгоритм',                    '2.3. Алгоритм'),
    ('2.5.',                             '2.5.'),
    ('2.6. Результаты',                  '2.6. Результаты'),
    ('2.6. Выводы',                      '2.6. Выводы'),
    ('3.1.',                             '3.1.'),
    ('3.2.',                             '3.2.'),
    ('3.3.',                             '3.3.'),
    ('3.4.',                             '3.4.'),
    ('3.5.',                             '3.5.'),
    ('ЗАКЛЮЧЕНИЕ',                       'ЗАКЛЮЧЕНИЕ'),
]

def find_rewrite_key(keyword, sections):
    """Find the first rewrite section key that contains the given keyword."""
    for k in sections:
        if keyword in k:
            return k
    return None

# ── Build heading index: {para_index: heading_text} ──────────────────────────
paras = doc.paragraphs
heading_indices = []  # list of (para_idx, heading_text)
for i, p in enumerate(paras):
    if is_heading(p):
        heading_indices.append((i, p.text.strip()))

print(f"\nFound {len(heading_indices)} heading paragraphs in original")

# ── Build section ranges: (start_para_idx, end_para_idx, heading_text) ────────
# start_para_idx = index of the heading paragraph itself
# end_para_idx   = index of the NEXT heading (exclusive), or len(paras)
section_ranges = []
for j, (idx, htxt) in enumerate(heading_indices):
    end = heading_indices[j + 1][0] if j + 1 < len(heading_indices) else len(paras)
    section_ranges.append((idx, end, htxt))

# ── Perform replacements (process in REVERSE order to keep indices valid) ─────
body = doc.element.body

def body_index(elem):
    """Return the index of elem in doc.element.body (top-level only)."""
    for i, child in enumerate(body):
        if child is elem:
            return i
    return -1

replacements_done = 0

for orig_kw, rw_kw in reversed(REPLACE_MAP):
    # Find matching section range in original
    matched_range = None
    for (start, end, htxt) in section_ranges:
        if orig_kw in htxt:
            matched_range = (start, end, htxt)
            break
    if matched_range is None:
        print(f"  WARNING: original heading containing '{orig_kw}' not found")
        continue

    start, end, htxt = matched_range

    # Find matching rewrite section
    rw_key = find_rewrite_key(rw_kw, rewrite_sections)
    if rw_key is None:
        print(f"  WARNING: rewrite section containing '{rw_kw}' not found")
        continue

    new_lines = rewrite_sections[rw_key]
    if not new_lines:
        print(f"  SKIP (empty): '{rw_key[:60]}'")
        continue

    # Collect paragraphs in this section to DELETE
    to_delete = []
    for i in range(start + 1, end):   # skip the heading itself
        p = paras[i]
        if not should_keep(p):
            to_delete.append(p)

    # DELETE old text paragraphs
    for p in to_delete:
        p._element.getparent().remove(p._element)

    # INSERT new paragraphs right after the heading
    heading_elem = paras[start]._element
    heading_body_idx = body_index(heading_elem)
    if heading_body_idx == -1:
        print(f"  ERROR: heading not found in body for '{htxt[:60]}'")
        continue

    # Insert in REVERSE order so they end up in correct order after the heading
    for line in reversed(new_lines):
        new_p = make_para_elem(line, template_para)
        # Insert immediately after heading
        heading_elem.addnext(new_p)

    print(f"  REPLACED [{len(to_delete):3d} del, {len(new_lines):3d} ins] '{htxt[:55]}'")
    replacements_done += 1

print(f"\nDone. {replacements_done} sections replaced.")

# ── Save ──────────────────────────────────────────────────────────────────────
out_path = '02_VKR_Ali_final.docx'
doc.save(out_path)
print(f"Saved → {out_path}")
